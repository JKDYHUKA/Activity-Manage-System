from django.shortcuts import render
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json
import jwt
import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from .models import *
from activities_organization.models import *
from activities_organization.activity_utils.organization_utils import convert_to_beijing_time
from chat.chat_utils import ask_openai


@csrf_exempt
def subQuestionnaire(request):
    if request.method == 'POST':
        body = json.loads(request.body)
        role=body['usertype']
        satisfaction=body['rate']
        suggestion=body['suggestion']
        notice_id=body['notice_id']
        notice=Notice.objects.get(id=notice_id)
        notice_num=Notice.objects.filter(activity_id=notice.activity_id).count()
        # 人数
        Part_num=ActivityParticipator.objects.filter(activity_id=notice.activity_id).count()
        activity=CreateActivity.objects.get(activity_id=notice.activity_id)
        people_num=Part_num+ActivityGuest.objects.filter(activity=activity).count()+1
        #互动数
        Feedback.objects.create(
                role=role,
                satisfaction=satisfaction,
                suggestions=suggestion,
                activity_id=notice.activity_id,
            )
        try:
            data = ActivityData.objects.get(activity_id=notice.activity_id)
            data.people_num=people_num
            data.notice_num=notice_num
            data.save()
        except ActivityData.DoesNotExist:
            ActivityData.objects.create(activity_id=activity.activity_id,
                                        people_num=people_num,
                                        notice_num=notice_num,)

        return JsonResponse({"message": "feedback successfully", 'code': '0'},
                                status=200)

    else:
        return JsonResponse({"message": "method error"}, status=401)


@csrf_exempt
def generate_report(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        act_id = data['act_id']

        act = CreateActivity.objects.get(activity_id=act_id)
        if act.activity_condition < 3:
            return JsonResponse({"message": "can not generate report now"}, status=200)
        act_data = ActivityData.objects.get(activity_id=act_id)
        feedback = Feedback.objects.filter(activity_id=act_id)

        act_name = act.activity_name
        act_description = act.activity_description
        act_leader = act.activity_leader.username
        start_time = act.activity_time.start_time
        end_time = act.activity_time.end_time
        act_start_time = convert_to_beijing_time(start_time)
        act_end_time = convert_to_beijing_time(end_time, is_end=True)
        act_type = act.activity_type
        act_person_number = act_data.people_num
        act_notice_num = act_data.notice_num
        act_chat_num = act_data.chat_num

        opinion = ""
        for f in feedback:
            role = f.role
            rate = f.satisfaction
            suggestion = f.suggestions
            opinion_content = f"[角色: {role}, 分数: {rate}, 建议: {suggestion}] "
            opinion = opinion + opinion_content

        content = f"活动名称： {act_name}, 活动描述: {act_description}, 活动创建者的用户名: {act_leader}, " \
                  f"活动起始时间: {act_start_time}, 活动结束时间: {act_end_time}, 活动类型: {act_type}, " \
                  f"活动参与人数: {act_person_number}, 活动发送的通知数量: {act_notice_num}, " \
                  f"活动聊天室中聊天信息数量: {act_chat_num}。\n\n以下是活动结束后收到的评价，每一条评价的信息都在中括号内"

        content = content + opinion
        system_content = "请生成一份中文活动报告，内容用大写汉字编号和三个自然段表示，不要有编号，绝对只能有且仅有三个自然段" \
                         "第一段根据活动描述和其他信息介绍该活动，第二段总结一下评价中的意见，最后一段请给出总结，注意不要直接重复活动描述"

        input_content = [{
                'role': 'system',
                'content': system_content,
            },
            {
                'role': 'user',
                'content': content
            }]

        response = ask_openai(input_content)
        print(response)

        doc = Document()

        doc.styles['Normal'].font.name = u'Microsoft JhengHei'

        title = f"\"{act_name}\"活动报告\n"
        title_paragraph = doc.add_heading(title, level=1)
        title_run = title_paragraph.runs[0]
        title_run.font.name = 'Microsoft JhengHei'
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        paragraphs = response.split('\n\n')
        print(paragraphs)

        num = ['一', '二', '三']
        sub_title = ['活动介绍', '活动评价', '活动总结']

        for i, para in enumerate(paragraphs, 1):
            # 添加小标题
            p_title = doc.add_paragraph()
            p_title.add_run(f'{num[i-1]}、 {sub_title[i-1]}').bold = True
            p_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 小标题左对齐

            # 添加段落内容
            p_content = doc.add_paragraph(para)
            p_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 段落内容左对齐
            p_content.paragraph_format.left_indent = Inches(0.5)  # 段落内容缩进两格

        base_path = os.getenv('BASED_PATH')
        doc_name = f"{base_path}//{str(act_id)}_{act_name}活动报告.docx"
        doc.save(doc_name)

        return JsonResponse({"message": "successfully"}, status=200)



